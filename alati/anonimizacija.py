import re
import os
from docx import Document

def anonimiziraj_tekst(tekst):
    """
    Glavna funkcija koja prepoznaje i mijenja osjetljive podatke u tekstu.
    """
    if not tekst:
        return ""

    # 1. JMBG (13 cifara u kontinuitetu)
    tekst = re.sub(r'\b\d{13}\b', '[ANONIMIZOVAN_JMBG]', tekst)
    
    # 2. Broj lične karte / pasoša (Npr. 00AAA0000 ili 9-cifreni brojevi)
    tekst = re.sub(r'\b\d{2}[A-Za-z]{3}\d{4}\b', '[ANONIMIZOVAN_BROJ_LK]', tekst)
    
    # 3. Brojevi telefona (Fiksni i mobilni u različitim formatima)
    telefon_pattern = r'(\+387|\+381|00387|00381|0)[6345][0-9][-\s./]*[0-9]{3}[-\s./]*[0-9]{3,4}'
    tekst = re.sub(telefon_pattern, '[ANONIMIZOVAN_TELEFON]', tekst)
    
    # 4. Email adrese
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    tekst = re.sub(email_pattern, '[ANONIMIZOVAN_EMAIL]', tekst)

    # 5. Ključne birokratske riječi za adrese (Ulica, ul., trg, naselje...)
    # Briše naziv ulice i broj koji slijede nakon ključne riječi
    adresa_pattern = r'(?i)\b(ulica|ul\.|trg|naselje|adresa|stanuje u)\b\s+([A-ZŽĆČŠĐ][a-zžćčšđ\s]+)+\s+\d+[a-zA-Z]?'
    tekst = re.sub(adresa_pattern, r'\1 [ANONIMIZOVANA_ADRESA]', tekst)

    return tekst

def obradi_fajl(putanja_ulaz):
    """
    Detektuje format fajla i vrši anonimizaciju.
    """
    ekstenzija = os.path.splitext(putanja_ulaz)[1].lower()
    putanja_izlaz = putanja_ulaz.replace(ekstenzija, f"_anonimizovano{ekstenzija}")

    if ekstenzija == '.txt':
        with open(putanja_ulaz, 'r', encoding='utf-8') as f:
            sadrzaj = f.read()
        
        cist_sadrzaj = anonimiziraj_tekst(sadrzaj)
        
        with open(putanja_izlaz, 'w', encoding='utf-8') as f:
            f.write(cist_sadrzaj)
        print(f"✅ Uspješno očišćeno! Rezultat sačuvan u: {putanja_izlaz}")

    elif ekstenzija == '.docx':
        doc = Document(putanja_ulaz)
        
        # Čišćenje paragrafa
        for p in doc.paragraphs:
            if p.text:
                p.text = anonimiziraj_tekst(p.text)
                
        # Čišćenje tabela ako postoje u dokumentu
        for tabela in doc.tables:
            for red in tabela.rows:
                for celija in red.cells:
                    if celija.text:
                        celija.text = anonimiziraj_tekst(celija.text)
                        
        doc.save(putanja_izlaz)
        print(f"✅ Word dokument očišćen! Rezultat sačuvan u: {putanja_izlaz}")
    else:
        print("❌ Format nije podržan. Molimo koristite .txt ili .docx fajlove.")

if __name__ == "__main__":
    print("--- Alat za građansku forenziku: Anonimizacija dokumenata ---")
    putanja = input("Unesite putanju do .txt ili .docx fajla (ili ga prevucite u terminal): ").strip("'\" ")
    if os.path.exists(putanja):
        obradi_fajl(putanja)
    else:
        print("❌ Fajl ne postoji na navedenoj putanji.")
